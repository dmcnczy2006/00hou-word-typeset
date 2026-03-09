"""
智能 Word 排版 Agent 主入口

提供 process_document 作为主入口函数，完成：解析用户意图 → 调度排版引擎 → 保存文档。
支持 Y Mode（单次再排版）和 Z Mode（多步骤工作流）。
"""

import logging
import re
from pathlib import Path
from typing import Literal, Optional

from dotenv import load_dotenv

load_dotenv()  # 加载 .env 中的 OPENAI_API_KEY 等配置

logger = logging.getLogger(__name__)

from .dispatcher.command import CommandDispatcher
from .dispatcher.workflow import WorkflowDispatcher
from .formatter.engine import WordProcessor
from .intent.parser import IntentParser
from .intent.workflow_parser import WorkflowParser


def _detect_z_mode(user_prompt: str) -> bool:
    """
    根据用户指令自动判断是否使用 Z Mode（多步骤工作流）

    触发条件：包含编号（1. 2. 3.）、分号分隔多句、或显式「步骤」等。
    """
    if not user_prompt or not user_prompt.strip():
        return False
    text = user_prompt.strip()
    if re.search(r"[1-9][\.\．]\s*\S", text):
        return True
    if "；" in text or ";" in text:
        parts = re.split(r"[；;]", text)
        if len(parts) >= 2:
            return True
    if "步骤" in text or "工作流" in text:
        return True
    return False


def process_document(
    file_path: str,
    user_prompt: str,
    preset: str = "official",
    output_path: Optional[str] = None,
    llm_connector=None,
    mode: Literal["y", "z", "auto"] = "auto",
) -> str:
    """
    主入口：解析用户意图 → 调度排版引擎 → 保存文档

    Args:
        file_path: 待处理的 .docx 文件路径
        user_prompt: 用户自然语言排版指令。Y Mode 如「把正文设为小四号仿宋」；
            Z Mode 如「1. 将标题2变为标题1；2. 设置多级列表；3. 再排版，正文小四」
        preset: 预设规则名称，如 "official"、"thesis"、"default"
        output_path: 输出文件路径，若为 None 则覆盖原文件
        llm_connector: LLM 连接器实例，若为 None 则使用 OpenAIConnector（从 .env 读取 API Key）
        mode: 工作模式。y=单次再排版，z=多步骤工作流，auto=根据指令自动判断

    Returns:
        处理后的文件路径

    Example:
        >>> process_document("report.docx", "正文小四仿宋，首行缩进2字符", preset="official")
        'report.docx'
    """
    # 1. 加载文档
    path = Path(file_path)
    logger.info("开始处理文档: %s", file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    if path.suffix.lower() != ".docx":
        raise ValueError("仅支持 .docx 格式")

    word_processor = WordProcessor.from_path(str(path))
    logger.info("文档加载成功，共 %d 个段落", len(word_processor.document.paragraphs))

    # 2. 模式判断与执行
    use_z_mode = mode == "z" or (mode == "auto" and _detect_z_mode(user_prompt))
    prompt_preview = user_prompt[:50] + "..." if len(user_prompt) > 50 else user_prompt
    logger.info("模式=%s，使用 Z Mode=%s，指令=%s", mode, use_z_mode, prompt_preview)

    if use_z_mode:
        workflow_parser = WorkflowParser(llm_connector=llm_connector)
        workflow = workflow_parser.parse(user_prompt)
        logger.info("工作流解析完成，共 %d 步", len(workflow.steps))
        WorkflowDispatcher.execute(
            workflow=workflow,
            word_processor=word_processor,
            llm_connector=llm_connector,
        )
    else:
        parser = IntentParser(llm_connector=llm_connector)
        intent = parser.parse(
            user_prompt=user_prompt,
            preset=preset,
            merge_with_preset=True,
        )
        logger.info("意图解析完成")
        CommandDispatcher.dispatch(intent, word_processor)

    # 4. 保存
    save_path = output_path or str(path)
    word_processor.save(save_path)
    logger.info("文档已保存至: %s", save_path)

    return save_path


def create_sample_document(path: str = "sample_00houTypeset.docx") -> str:
    """
    创建示例文档（用于测试）

    Args:
        path: 保存路径

    Returns:
        保存后的文件路径
    """
    from docx import Document

    doc = Document()
    doc.add_heading("示例文档", 0)
    doc.add_paragraph(
        "这是一段正文。智能 Word 排版 Agent 可以根据自然语言指令自动调整字体、字号、段落格式等。"
    )
    doc.add_paragraph(
        "例如：「把正文设为小四号仿宋，首行缩进2字符」即可一键应用公文标准格式。"
    )
    doc.save(path)
    return path


if __name__ == "__main__":
    import argparse
    import sys

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%H:%M:%S",
    )

    parser = argparse.ArgumentParser(
        description="智能 Word 排版 Agent：Y Mode 单次再排版，Z Mode 多步骤工作流",
    )
    parser.add_argument("docx", help="待处理的 .docx 文件路径")
    parser.add_argument(
        "prompt",
        nargs="?",
        default="",
        help="用户自然语言指令。Z Mode 示例：1. 将标题2变为标题1；2. 设置多级列表；3. 再排版，正文小四",
    )
    parser.add_argument(
        "-m",
        "--mode",
        choices=["y", "z", "auto"],
        default="auto",
        help="工作模式：y=单次再排版，z=多步骤工作流，auto=自动判断（默认）",
    )
    parser.add_argument(
        "-p",
        "--preset",
        default="official",
        help="预设规则名称（默认 official）",
    )
    parser.add_argument(
        "-o",
        "--output",
        default=None,
        help="输出文件路径，默认覆盖原文件",
    )
    args = parser.parse_args()

    doc_path = args.docx
    prompt = args.prompt or ""

    if not Path(doc_path).exists():
        create_sample_document(doc_path)
        print(f"已创建示例文档: {doc_path}")

    output_path = args.output or doc_path
    result = process_document(
        doc_path,
        prompt,
        preset=args.preset,
        output_path=output_path,
        mode=args.mode,
    )
    print(f"处理完成，已保存至: {result}")
