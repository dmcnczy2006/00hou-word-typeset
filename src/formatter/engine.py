"""
Word 排版引擎

封装 python-docx 操作，提供字体、段落、文本替换等排版原子能力。
"""

import logging
import re
from typing import Optional

logger = logging.getLogger(__name__)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ..schemas.typesetting import FontConfig, ParagraphConfig


# 中文字体名到 Word 内部名称的映射（智能样式映射）
FONT_NAME_MAP = {
    "仿宋": "仿宋",
    "仿宋_GB2312": "仿宋_GB2312",
    "宋体": "宋体",
    "黑体": "黑体",
    "楷体": "楷体",
    "微软雅黑": "微软雅黑",
    "Times New Roman": "Times New Roman",
    "Arial": "Arial",
}


def _parse_color(color_str: Optional[str]):
    """
    解析颜色字符串为 RGBColor

    支持 #RRGGBB 或 #RGB 格式，以及常见颜色名称。
    """
    if not color_str:
        return None
    from docx.shared import RGBColor

    color_str = color_str.strip()
    # Hex 格式
    if color_str.startswith("#"):
        hex_val = color_str[1:]
        if len(hex_val) == 6:
            r = int(hex_val[0:2], 16)
            g = int(hex_val[2:4], 16)
            b = int(hex_val[4:6], 16)
            return RGBColor(r, g, b)
        elif len(hex_val) == 3:
            r = int(hex_val[0] * 2, 16)
            g = int(hex_val[1] * 2, 16)
            b = int(hex_val[2] * 2, 16)
            return RGBColor(r, g, b)
    # 常见颜色名称
    color_names = {
        "black": (0, 0, 0),
        "white": (255, 255, 255),
        "red": (255, 0, 0),
        "green": (0, 128, 0),
        "blue": (0, 0, 255),
    }
    lower = color_str.lower()
    if lower in color_names:
        r, g, b = color_names[lower]
        return RGBColor(r, g, b)
    return None


def _alignment_from_str(align: Optional[str]):
    """将字符串对齐方式转换为 WD_ALIGN_PARAGRAPH 枚举"""
    if not align:
        return None
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align.lower())


# target_scope -> Word 内置样式名（python-docx 使用英文名）
SCOPE_TO_STYLE_NAMES: dict[str, list[str]] = {
    "body": ["Normal"],
    "heading": ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6", "Heading 7", "Heading 8", "Heading 9"],
    "heading_1": ["Heading 1"],
    "heading_2": ["Heading 2"],
    "heading_3": ["Heading 3"],
    "heading_4": ["Heading 4"],
    "heading_5": ["Heading 5"],
    "heading_6": ["Heading 6"],
    "heading_7": ["Heading 7"],
    "heading_8": ["Heading 8"],
    "heading_9": ["Heading 9"],
    "caption": ["Caption"],
    "all": ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6", "Heading 7", "Heading 8", "Heading 9", "Caption"],
}


def _get_style_names_for_scope(target_scope: str) -> list[str]:
    """根据 target_scope 返回要修改的 Word 样式名列表"""
    return SCOPE_TO_STYLE_NAMES.get(target_scope, ["Normal"])


def _resolve_font(raw: Optional[str]) -> Optional[str]:
    """解析字体名（智能映射）"""
    if not raw:
        return None
    return FONT_NAME_MAP.get(raw, raw)


# 主题字体属性：在 Heading 等样式中会覆盖显式字体，需移除后显式字体才生效
# 注：OOXML 中为 cstheme（小写）
_THEME_FONT_ATTRS = (
    qn("w:asciiTheme"),
    qn("w:eastAsiaTheme"),
    qn("w:hAnsiTheme"),
    qn("w:cstheme"),
)


def _remove_theme_font_attrs(rFonts) -> None:
    """移除 rFonts 中的主题字体属性，使显式字体（ascii/eastAsia 等）生效"""
    for attr in _THEME_FONT_ATTRS:
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]


def _apply_font_to_style(style, font_config: FontConfig) -> None:
    """
    将 FontConfig 应用到 Word 样式定义（修改样式本身，而非直接格式）
    需同时设置 ascii 和 eastAsia 以支持中文字体。
    Heading 等样式使用主题字体，会覆盖显式字体，故需先移除主题属性。
    """
    font_ascii = _resolve_font(font_config.name_ascii or font_config.name)
    font_east_asia = _resolve_font(font_config.name_east_asia or font_config.name)
    rgb_color = _parse_color(font_config.color)

    font = style.font
    if font_ascii or font_east_asia:
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        _remove_theme_font_attrs(rFonts)  # 移除主题字体，使显式字体生效
    if font_ascii:
        font.name = font_ascii
    if font_east_asia:
        try:
            rPr = style._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), font_east_asia)
        except Exception:
            font.name = font_east_asia
    if font_config.size_pt is not None:
        font.size = Pt(font_config.size_pt)
    if rgb_color:
        font.color.rgb = rgb_color
    if font_config.bold is not None:
        font.bold = font_config.bold
    if font_config.italic is not None:
        font.italic = font_config.italic
    if font_config.underline is not None:
        font.underline = font_config.underline


def _apply_font_to_run(run, font_config: FontConfig) -> None:
    """
    将 FontConfig 应用到 run 的直接格式（批量设置已有内容的字体）
    """
    font_ascii = _resolve_font(font_config.name_ascii or font_config.name)
    font_east_asia = _resolve_font(font_config.name_east_asia or font_config.name)
    rgb_color = _parse_color(font_config.color)

    if font_ascii:
        run.font.name = font_ascii
    if font_east_asia:
        try:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), font_east_asia)
        except Exception:
            run.font.name = font_east_asia
    if font_config.size_pt is not None:
        run.font.size = Pt(font_config.size_pt)
    if rgb_color:
        run.font.color.rgb = rgb_color
    if font_config.bold is not None:
        run.font.bold = font_config.bold
    if font_config.italic is not None:
        run.font.italic = font_config.italic
    if font_config.underline is not None:
        run.font.underline = font_config.underline


class WordProcessor:
    """
    Word 文档排版处理器

    封装 python-docx 的文档操作，提供批量样式设置、文本替换、段落整形等功能。
    """

    def __init__(self, document: Document):
        """
        初始化排版处理器

        Args:
            document: python-docx Document 对象
        """
        self._doc = document

    @classmethod
    def from_path(cls, path: str) -> "WordProcessor":
        """
        从文件路径加载文档并创建 WordProcessor

        Args:
            path: .docx 文件路径

        Returns:
            WordProcessor 实例
        """
        logger.debug("加载文档: %s", path)
        doc = Document(path)
        return cls(doc)

    def save(self, path: Optional[str] = None) -> str:
        """
        保存文档

        Args:
            path: 保存路径，若为 None 则覆盖原文件（需在 from_path 加载时有效）

        Returns:
            保存后的文件路径
        """
        if path:
            logger.debug("保存文档至: %s", path)
            self._doc.save(path)
            return path
        raise ValueError("保存路径不能为空")

    def set_global_styles(
        self,
        font_config: FontConfig,
        target_scope: str = "all",
    ) -> None:
        """
        同时执行：1) 修改 Word 文档的样式定义；2) 批量设置已有内容的字体直接格式。

        - 样式定义：Normal、Heading 1 等会更新，新输入内容将自动继承
        - 批量设置：遍历匹配范围的段落/runs，直接应用字体格式，确保已有内容立即生效

        Args:
            font_config: 字体配置
            target_scope: 作用范围。支持：body、heading、heading_1~heading_9、caption、all
        """
        # 1. 修改样式定义
        style_names = _get_style_names_for_scope(target_scope)
        logger.info(
            "set_global_styles: 修改样式定义，范围=%s, 样式=%s",
            target_scope,
            style_names,
        )
        for style_name in style_names:
            try:
                style = self._doc.styles[style_name]
                _apply_font_to_style(style, font_config)
            except KeyError:
                logger.debug("样式 %s 不存在于文档中，跳过", style_name)

        # 2. 批量设置已有内容的字体直接格式
        run_count = 0
        for para in self._doc.paragraphs:
            if not self._match_scope(para, target_scope):
                continue
            for run in para.runs:
                _apply_font_to_run(run, font_config)
                run_count += 1
        logger.info("set_global_styles 完成：样式定义已更新，共处理 %d 个 run", run_count)

    def semantic_replace(
        self,
        search: str,
        replace: str,
        use_regex: bool = False,
    ) -> int:
        """
        智能文本替换（支持正则表达式）

        Args:
            search: 查找内容
            replace: 替换内容
            use_regex: 是否将 search 视为正则表达式

        Returns:
            替换次数
        """
        count = 0
        for para in self._doc.paragraphs:
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                if use_regex:
                    new_text, n = re.subn(search, replace, text)
                else:
                    new_text = text.replace(search, replace)
                    n = text.count(search)
                if n > 0:
                    run.text = new_text
                    count += n

        # 同时处理表格中的文本
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text = run.text
                            if not text:
                                continue
                            if use_regex:
                                new_text, n = re.subn(search, replace, text)
                            else:
                                new_text = text.replace(search, replace)
                                n = text.count(search)
                            if n > 0:
                                run.text = new_text
                                count += n

        logger.info("semantic_replace 完成: %r -> %r, 共替换 %d 处", search, replace, count)
        return count

    def paragraph_shaping(
        self,
        paragraph_config: ParagraphConfig,
        target_scope: str = "all",
    ) -> None:
        """
        同时执行：1) 修改 Word 文档样式的段落格式定义；2) 批量设置已有段落的直接格式。

        - 样式定义：新输入的内容将自动继承行距、首行缩进等
        - 批量设置：遍历匹配范围的段落，直接应用段落格式，确保已有内容立即生效

        Args:
            paragraph_config: 段落配置
            target_scope: 作用范围
        """
        # 1. 修改样式定义
        style_names = _get_style_names_for_scope(target_scope)
        logger.info(
            "paragraph_shaping: 修改样式定义，范围=%s, 样式=%s, 首行缩进=%s, 行距=%s",
            target_scope,
            style_names,
            paragraph_config.first_line_indent,
            paragraph_config.line_spacing,
        )
        for style_name in style_names:
            try:
                style = self._doc.styles[style_name]
                pf = style.paragraph_format
                if paragraph_config.first_line_indent is not None:
                    pf.first_line_indent = Pt(paragraph_config.first_line_indent)
                if paragraph_config.line_spacing is not None:
                    pf.line_spacing = paragraph_config.line_spacing
                if paragraph_config.space_before is not None:
                    pf.space_before = Pt(paragraph_config.space_before)
                if paragraph_config.space_after is not None:
                    pf.space_after = Pt(paragraph_config.space_after)
                align = _alignment_from_str(paragraph_config.alignment)
                if align is not None:
                    pf.alignment = align
            except KeyError:
                logger.debug("样式 %s 不存在于文档中，跳过", style_name)

        # 2. 批量设置已有段落的直接格式
        para_count = 0
        for para in self._doc.paragraphs:
            if not self._match_scope(para, target_scope):
                continue
            pf = para.paragraph_format
            if paragraph_config.first_line_indent is not None:
                pf.first_line_indent = Pt(paragraph_config.first_line_indent)
            if paragraph_config.line_spacing is not None:
                pf.line_spacing = paragraph_config.line_spacing
            if paragraph_config.space_before is not None:
                pf.space_before = Pt(paragraph_config.space_before)
            if paragraph_config.space_after is not None:
                pf.space_after = Pt(paragraph_config.space_after)
            align = _alignment_from_str(paragraph_config.alignment)
            if align is not None:
                pf.alignment = align
            para_count += 1
        logger.info("paragraph_shaping 完成：样式定义已更新，共处理 %d 个段落", para_count)

    def apply_first_line_indent(
        self,
        indent_pt: float,
        target_scope: str = "body",
    ) -> None:
        """
        修改样式的首行缩进定义

        Args:
            indent_pt: 缩进量（磅），2 字符 ≈ 24pt
            target_scope: 作用范围（body、heading、heading_1~9、caption、all）
        """
        self.paragraph_shaping(
            ParagraphConfig(first_line_indent=indent_pt),
            target_scope=target_scope,
        )

    def change_heading_level(self, from_scope: str, to_scope: str) -> int:
        """
        将指定级别的标题改为另一级别（如 heading_2 → heading_1）

        Args:
            from_scope: 源 scope，如 heading_2
            to_scope: 目标 scope，如 heading_1

        Returns:
            修改的段落数量
        """
        from_style = SCOPE_TO_STYLE_NAMES.get(from_scope, [None])[0]
        to_style = SCOPE_TO_STYLE_NAMES.get(to_scope, [None])[0]
        if not from_style or not to_style:
            logger.warning("change_heading_level: 无效 scope %s 或 %s", from_scope, to_scope)
            return 0
        count = 0
        for para in self._doc.paragraphs:
            if self._match_scope(para, from_scope):
                para.style = self._doc.styles[to_style]
                count += 1
        logger.info("change_heading_level: %s -> %s, 共修改 %d 个段落", from_scope, to_scope, count)
        return count

    def apply_multilevel_list(self, target_scope: str, list_style: str = "multilevel") -> int:
        """
        为匹配的段落应用多级列表编号

        MVP：仅支持 target_scope: heading，对 Heading 1~9 按层级设置 ilvl（0~8）。
        使用文档中已有的 numbering 定义（numId=1）。

        Args:
            target_scope: 作用范围，heading 表示仅标题
            list_style: 列表样式（multilevel/decimal/bullet），当前均使用多级编号

        Returns:
            应用的段落数量
        """
        if target_scope == "heading":
            return self._apply_list_to_headings()
        if target_scope == "body":
            return self._apply_list_to_body()
        if target_scope == "all":
            return self._apply_list_to_headings() + self._apply_list_to_body()
        logger.warning("apply_multilevel_list: 不支持的 target_scope %s", target_scope)
        return 0

    def _apply_list_to_headings(self) -> int:
        """为标题段落应用多级列表编号"""
        num_id = 1
        count = 0
        for para in self._doc.paragraphs:
            if not self._is_heading(para):
                continue
            level = self._get_heading_level(para)
            if level is None:
                continue
            ilvl = min(level - 1, 8)
            self._set_paragraph_num_pr(para, ilvl, num_id)
            count += 1
        logger.info("apply_multilevel_list: 标题范围，共应用 %d 个段落", count)
        return count

    def _apply_list_to_body(self) -> int:
        """为正文段落应用多级列表编号（ilvl=0）"""
        num_id = 1
        count = 0
        for para in self._doc.paragraphs:
            if not self._match_scope(para, "body"):
                continue
            self._set_paragraph_num_pr(para, 0, num_id)
            count += 1
        logger.info("apply_multilevel_list: 正文范围，共应用 %d 个段落", count)
        return count

    def _set_paragraph_num_pr(self, paragraph, ilvl: int, num_id: int) -> None:
        """为段落设置 numPr（ilvl + numId）"""
        p_pr = paragraph._p.get_or_add_pPr()
        old_num_pr = p_pr.find(qn("w:numPr"))
        if old_num_pr is not None:
            p_pr.remove(old_num_pr)
        num_pr = OxmlElement("w:numPr")
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), str(ilvl))
        num_id_elem = OxmlElement("w:numId")
        num_id_elem.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl_elem)
        num_pr.append(num_id_elem)
        p_pr.append(num_pr)

    def set_page_margins(
        self,
        top_pt: Optional[float] = None,
        bottom_pt: Optional[float] = None,
        left_pt: Optional[float] = None,
        right_pt: Optional[float] = None,
    ) -> None:
        """设置页面边距（磅）。未指定的边距保持不变。"""
        for section in self._doc.sections:
            if top_pt is not None:
                section.top_margin = Pt(top_pt)
            if bottom_pt is not None:
                section.bottom_margin = Pt(bottom_pt)
            if left_pt is not None:
                section.left_margin = Pt(left_pt)
            if right_pt is not None:
                section.right_margin = Pt(right_pt)
        logger.info(
            "set_page_margins: top=%s, bottom=%s, left=%s, right=%s",
            top_pt, bottom_pt, left_pt, right_pt,
        )

    def set_header_footer(
        self,
        header_text: Optional[str] = None,
        footer_text: Optional[str] = None,
        odd_even: bool = False,
        odd_header_text: Optional[str] = None,
        even_header_text: Optional[str] = None,
        odd_footer_text: Optional[str] = None,
        even_footer_text: Optional[str] = None,
        show_chapter: bool = False,
        odd_show_chapter: Optional[bool] = None,
        even_show_chapter: Optional[bool] = None,
        underline: bool = False,
        odd_underline: Optional[bool] = None,
        even_underline: Optional[bool] = None,
        page_number: bool = False,
        odd_page_number: Optional[bool] = None,
        even_page_number: Optional[bool] = None,
        first_page_different: bool = False,
    ) -> None:
        """
        设置页眉页脚。odd_even=True 时分别设置奇数页和偶数页，可指定不同内容。
        """
        try:
            self._doc.settings.odd_and_even_pages_header_footer = odd_even
        except AttributeError:
            pass

        def _odd(val, fallback):
            return val if val is not None else fallback

        for section in self._doc.sections:
            section.different_first_page_header_footer = first_page_different

            if odd_even:
                h_odd = odd_header_text if odd_header_text is not None else header_text
                h_even = even_header_text if even_header_text is not None else header_text
                f_odd = odd_footer_text if odd_footer_text is not None else footer_text
                f_even = even_footer_text if even_footer_text is not None else footer_text
                sc_odd = _odd(odd_show_chapter, show_chapter)
                sc_even = _odd(even_show_chapter, show_chapter)
                ul_odd = _odd(odd_underline, underline)
                ul_even = _odd(even_underline, underline)
                pn_odd = _odd(odd_page_number, page_number)
                pn_even = _odd(even_page_number, page_number)

                if h_odd or sc_odd or ul_odd:
                    self._setup_header(section.header, h_odd, sc_odd, ul_odd)
                if f_odd or pn_odd:
                    self._setup_footer(section.footer, f_odd, pn_odd)
                if h_even or sc_even or ul_even:
                    self._setup_header(section.even_page_header, h_even, sc_even, ul_even)
                if f_even or pn_even:
                    self._setup_footer(section.even_page_footer, f_even, pn_even)
            else:
                if header_text or show_chapter or underline:
                    self._setup_header(section.header, header_text, show_chapter, underline)
                if footer_text or page_number:
                    self._setup_footer(section.footer, footer_text, page_number)

        logger.info(
            "set_header_footer: odd_even=%s, header=%s, footer=%s, page_number=%s",
            odd_even, header_text or odd_header_text, footer_text or odd_footer_text, page_number,
        )

    def _setup_header(self, header, text: Optional[str], show_chapter: bool, underline: bool) -> None:
        """配置页眉"""
        header.is_linked_to_previous = False
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        self._clear_paragraph(para)
        if show_chapter:
            self._add_styleref_field(para, "Heading 1")
        if text:
            para.add_run(text)
        if underline:
            self._add_border_bottom(para)

    def _setup_footer(self, footer, text: Optional[str], page_number: bool) -> None:
        """配置页脚"""
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        self._clear_paragraph(para)
        if text:
            para.add_run(text)
        if page_number:
            if text:
                para.add_run(" ")
            self._add_page_number_field(para)

    def _clear_paragraph(self, paragraph) -> None:
        """清空段落内容（保留段落属性）"""
        for child in list(paragraph._p):
            if child.tag == qn("w:r"):
                paragraph._p.remove(child)

    def _add_page_number_field(self, paragraph) -> None:
        """在段落中添加页码字段"""
        run = paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        fld_char3 = OxmlElement("w:fldChar")
        fld_char3.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr)
        run._r.append(fld_char2)
        run._r.append(fld_char3)

    def _add_styleref_field(self, paragraph, style_name: str) -> None:
        """在段落中添加 STYLEREF 字段（显示当前章节）"""
        run = paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' STYLEREF "{style_name}" \\n'
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        fld_char3 = OxmlElement("w:fldChar")
        fld_char3.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr)
        run._r.append(fld_char2)
        run._r.append(fld_char3)

    def _add_border_bottom(self, paragraph) -> None:
        """为段落添加下边框"""
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.get_or_add_pBdr()
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        p_bdr.append(bottom)

    def update_fields(self) -> None:
        """设置文档在打开时自动更新域（Word 打开时会提示更新）"""
        elem = self._doc.settings.element
        update_tag = qn("w:updateFields")
        existing = next((c for c in elem if c.tag == update_tag), None)
        if existing is None:
            update_elem = OxmlElement("w:updateFields")
            update_elem.set(qn("w:val"), "true")
            elem.append(update_elem)
        else:
            existing.set(qn("w:val"), "true")
        logger.info("update_fields: 已设置打开时更新域")

    def collapse_empty_lines(self, from_count: int, to_count: int) -> int:
        """
        将 x 个连续空行改成 y 个空行。

        Args:
            from_count: 连续空行数量 x（>=2）
            to_count: 目标空行数量 y（>=0）

        Returns:
            移除的段落数量
        """
        if from_count < 2:
            return 0
        paras = list(self._doc.paragraphs)
        to_remove = []
        i = 0
        while i < len(paras):
            if not self._is_empty_paragraph(paras[i]):
                i += 1
                continue
            empty_start = i
            while i < len(paras) and self._is_empty_paragraph(paras[i]):
                i += 1
            empty_len = i - empty_start
            if empty_len >= from_count:
                for j in range(to_count, empty_len):
                    to_remove.append(paras[empty_start + j]._p)
        for p_elem in reversed(to_remove):
            parent = p_elem.getparent()
            if parent is not None:
                parent.remove(p_elem)
        logger.info("collapse_empty_lines: %d 个连续空行->%d 个, 共移除 %d 段", from_count, to_count, len(to_remove))
        return len(to_remove)

    def _is_empty_paragraph(self, paragraph) -> bool:
        """判断段落是否为空（无内容或仅空白）"""
        return not paragraph.text.strip()

    def _is_heading(self, paragraph) -> bool:
        """判断段落是否为标题样式"""
        return self._get_heading_level(paragraph) is not None

    def _get_heading_level(self, paragraph) -> Optional[int]:
        """
        获取段落的标题级别（1-9），非标题返回 None
        """
        if paragraph.style is None:
            return None
        style_name = getattr(paragraph.style, "name", "") or ""
        if "heading" not in style_name.lower() and "标题" not in style_name:
            return None
        # 匹配 Heading 1, 标题 1, Heading1 等
        m = re.search(r"(?:heading|标题)\s*(\d)", style_name, re.IGNORECASE)
        if m:
            return int(m.group(1))
        return 1  # 未明确级别时默认 1

    def _is_caption(self, paragraph) -> bool:
        """判断段落是否为题注样式"""
        if paragraph.style is None:
            return False
        style_name = getattr(paragraph.style, "name", "") or ""
        name_lower = style_name.lower()
        return "caption" in name_lower or "题注" in style_name

    def _match_scope(self, paragraph, target_scope: str) -> bool:
        """
        判断段落是否匹配目标作用范围

        支持：body、heading、heading_1~heading_9、caption、all
        """
        if target_scope == "all":
            return True
        if target_scope == "body":
            return not self._is_heading(paragraph) and not self._is_caption(paragraph)
        if target_scope == "heading":
            return self._is_heading(paragraph)
        if target_scope.startswith("heading_"):
            try:
                level = int(target_scope.split("_")[1])
                return self._get_heading_level(paragraph) == level
            except (ValueError, IndexError):
                return False
        if target_scope == "caption":
            return self._is_caption(paragraph)
        if target_scope == "emphasis":
            # emphasis 为 run 级，段落级视为不匹配（由后续扩展处理）
            return False
        return True

    @property
    def document(self) -> Document:
        """获取底层 Document 对象"""
        return self._doc
